"""DocxParser — DOCX 文件解析器。

使用 python-docx 提取 Word 文档中的段落文本。
保留段落顺序，段落之间以换行分隔。
"""

from pathlib import Path

from document.model import Document, DocumentMetadata, DocumentType

from .base import BaseParser
from .exceptions import ParserError


class DocxParser(BaseParser):
    """DOCX 解析器，对应 DocumentType.DOCX。"""

    def parse(self, file_path: str) -> Document:
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")
        if not path.is_file():
            raise ParserError(f"路径不是文件: {file_path}")

        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ParserError(
                "python-docx 未安装，请执行: pip install python-docx"
            )

        try:
            doc = DocxDocument(str(path))
        except Exception as e:
            raise ParserError(f"无法打开 DOCX 文件: {file_path}") from e

        paragraphs: list[str] = []
        try:
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
        except Exception as e:
            raise ParserError(f"提取 DOCX 文本时出错: {file_path}") from e

        content = "\n".join(paragraphs)

        return Document(
            content=content,
            metadata=DocumentMetadata(
                filename=path.name,
                file_type=DocumentType.DOCX,
            ),
        )
